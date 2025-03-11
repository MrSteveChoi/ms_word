import os
import win32com.client

from docxcompose.composer import Composer
from docx import Document as Document_compose

from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def get_document_info(doc):
    """
    Word 문서의 주요 정보를 하나의 함수에서 추출합니다.
    """
    try:
        # 초기 데이터 구조
        document_info = {
            "total_height": None,
            "fonts": {},
            "spacing_info": [],
            "margins": {},
            "content_height": {
                "text_height": 0,
                "image_height": 0,
                "total_height": 0
            }
        }

        # 총 세로 길이 계산
        num_pages = doc.ComputeStatistics(2)  # 2: wdStatisticPages
        page_height = doc.PageSetup.PageHeight  # 포인트 단위
        document_info["total_height"] = page_height

        # 폰트 정보 추출
        fonts = {}
        for paragraph in doc.Paragraphs:
            font_name = paragraph.Range.Font.Name
            font_size = paragraph.Range.Font.Size
            if font_name not in fonts:
                fonts[font_name] = set()
            fonts[font_name].add(font_size)
        document_info["fonts"] = {font: list(sizes) for font, sizes in fonts.items()}

        # 자간 및 행간 정보 추출
        for paragraph in doc.Paragraphs:
            line_spacing = paragraph.LineSpacing
            for character in paragraph.Range.Characters:
                char_spacing = character.Font.Spacing
                document_info["spacing_info"].append({
                    "line_spacing": line_spacing,
                    "char_spacing": char_spacing
                })

        # 여백 정보 추출
        document_info["margins"] = {
            "top_margin": doc.PageSetup.TopMargin,
            "bottom_margin": doc.PageSetup.BottomMargin, 
            "header_dist": doc.PageSetup.HeaderDistance, 
            "footer_dist": doc.PageSetup.FooterDistance
        }

        # 내용의 전체 높이 계산
        total_text_height = 0
        total_image_height = 0

        for paragraph in doc.Paragraphs:
            font_size = paragraph.Range.Font.Size
            line_spacing = paragraph.LineSpacing
            line_count = paragraph.Range.ComputeStatistics(1)  # 1은 wdStatisticLines
            paragraph_height = (font_size * line_count) + (line_spacing * (line_count - 1))
            total_text_height += paragraph_height

        for shape in doc.InlineShapes:
            total_image_height += shape.Height

        document_info["content_height"] = {
            "text_height": total_text_height,
            "image_height": total_image_height,
            "total_height": total_text_height + total_image_height
        }

        return document_info
    
    except:
        print("Error")

    # finally:
    #     print("문서 정보 확인 완료")


def group_docs_by_page(base_path, file_list):
    """
    base_path : 단일 파일이 들어있는 폴더의 상위 경로 (ex: .../word_files/q)
    file_list : Word 파일 이름의 리스트 (ex: ["Q_short_01.docx", "Q_short_02.docx", ...])

    각 문서의 pt를 측정하여,
    한 페이지에 들어갈 수 있는 문서들을 한 덩어리(리스트)로 묶고,
    페이지가 넘어가면 새 리스트를 생성해 2차원 리스트를 반환합니다.

    수정 요청사항:
      "문제 + 한 줄 띄우기 + 문제" 의 높이를
      문제1 높이 + 행간 + 빈 줄 높이 + 행간 + 문제2 높이
      로 계산하여 배치가 가능하면 같은 페이지에, 불가능하면 페이지를 넘긴다.
    """

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False

    grouped_pages = []   # 최종 결과(2차원 리스트)
    current_page = []    # 현재 페이지에 들어갈 파일들
    current_text_pt = 0  # 현재 페이지에 누적된 문서 높이

    usable_pt = 0        # 한 페이지에서 쓸 수 있는 최대 pt (마스터 문서 기준)
    line_spacing_pt = 0  # 기본 행간(pt)
    
    # "빈 한 줄 높이(폰트 크기)"를 고정 12pt로 예시 설정
    # (get_document_info에서 실제 폰트 크기를 받아오면 그 값으로 대체 가능)
    blank_line_pt = 0

    for idx, filename in enumerate(file_list):
        doc_path = os.path.join(base_path, filename)
        
        # Word 문서 열기
        doc = word.Documents.Open(doc_path, Visible=False)
        result = get_document_info(doc)
        doc.Close(False)

        # 이 문서(문제)의 실제 내용 높이
        inner_text_pt = result["content_height"]["total_height"]

        # -------------------------------------
        # 1) 첫 파일(마스터 문서) 처리
        # -------------------------------------
        if idx == 0:
            # usable_pt 계산 = (문서 전체 높이) - (위/아래 마진)
            total_height = result["total_height"]
            top_margin = result['margins']["top_margin"]
            bottom_margin = result['margins']["bottom_margin"]
            header_dist = result["margins"]["header_dist"]
            footer_dist = result["margins"]["footer_dist"]
            usable_pt = total_height - (top_margin + bottom_margin + header_dist + footer_dist)
            blank_line_pt = list(result["fonts"].values())[-1][0]

            # 행간 (예: 첫 문서에서 가져온다)
            line_spacing_pt = result['spacing_info'][0]['line_spacing']

            # 첫 문서는 그냥 현재 페이지에 추가
            current_page.append(filename)
            current_text_pt = inner_text_pt

        # -------------------------------------
        # 2) 두 번째 문서부터
        # -------------------------------------
        else:
            # 만약 현재 페이지가 비어있지 않다면(이미 문제가 하나 이상 들어있다면)
            # "문제 + 한 줄 띄우기 + 문제"의 높이를 고려해야 함
            if len(current_page) > 0:
                # 추가로 필요한 높이 = 행간 + 빈줄 + 행간 + 다음 문서 높이
                needed_pt = line_spacing_pt + blank_line_pt + line_spacing_pt + inner_text_pt
            else:
                # 새 페이지에 바로 배치하는 경우라면, 빈 줄 없이 그냥 이 문서 높이만 추가
                needed_pt = inner_text_pt

            # 같은 페이지에 들어갈 수 있는지 판단
            if (current_text_pt + needed_pt) < usable_pt-200:
                # 같은 페이지에 추가
                current_page.append(filename)
                current_text_pt += needed_pt
            else:
                # 페이지를 넘겨야 함
                grouped_pages.append(current_page)
                # 새 페이지 초기화
                current_page = [filename]
                current_text_pt = inner_text_pt

    # 모든 문서 처리 후, 마지막 페이지가 비어있지 않다면 결과에 넣는다
    if current_page:
        grouped_pages.append(current_page)

    word.Quit()

    return grouped_pages


def combine_all_docx(full_dir, files_list, output_filename, is_reset_num):
    """
    주어진 file들의 list를 토대로 모든 문서를 결합.

    full_dir        : 파일들이 위치한 폴더의 절대경로
    files_list      : 병합할 파일들의 2차원 리스트
    output_filename : 유저가 입력한 최종 출력물의 이름
    is_reset_num    : 문제 번호 초기화 여부
    """
    # 1) 첫 번째 서브 리스트의 첫 번째 파일을 '마스터 문서'로 설정
    master = Document_compose(os.path.join(full_dir, files_list[0][0]))
    composer = Composer(master)

    
    # 문제 번호 reset을 위한 변수
    reset_num_count = 1

    ### ----------------- 문제 번호 reset -----------------
    if is_reset_num:
        if master.paragraphs:
            paragraph = master.paragraphs[0]
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # ✅ 왼쪽 정렬

            if paragraph.runs:
                paragraph.clear()  # ✅ 기존 번호 삭제
                new_run = paragraph.add_run(f"{reset_num_count}. ")  # ✅ 새로운 번호 추가
                new_run.bold = True  # ✅ Bold 적용
                new_run.font.size = Pt(12)  # ✅ 글자 크기 설정

        reset_num_count += 1  # ✅ 문제 번호 증가
    ### ----------------- 문제 번호 reset -----------------

    # 만약 첫 번째 서브 리스트에 파일이 여러 개라면, 나머지 파일을 같은 페이지에 순차적으로 추가
    first_sub_list = files_list[0]
    if len(first_sub_list) > 1:
        # 두 번째 파일부터 병합
        for file_name in first_sub_list[1:]:
            composer.doc.add_paragraph()  # 한 줄 띄우고
            # doc_temp = Document_compose(os.path.join(full_dir, f"{file_name}.docx"))
            doc_temp = Document_compose(os.path.join(full_dir, file_name))
            
            ### ----------------- 문제 번호 reset -----------------
            if is_reset_num:
                if doc_temp.paragraphs:
                    paragraph = doc_temp.paragraphs[0]
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # ✅ 왼쪽 정렬

                    if paragraph.runs:
                        paragraph.clear()  # ✅ 기존 번호 삭제
                        new_run = paragraph.add_run(f"{reset_num_count}. ")  # ✅ 새로운 번호 추가
                        new_run.bold = True  # ✅ Bold 적용
                        new_run.font.size = Pt(12)  # ✅ 글자 크기 설정

                reset_num_count += 1  # ✅ 문제 번호 증가
            ### ----------------- 문제 번호 reset -----------------

            composer.append(doc_temp)

   

    # 2) 나머지 서브 리스트(1번째 인덱스부터 끝까지) 처리
    for sub_list_index in range(1, len(files_list)):
        # 새로운 서브 리스트를 삽입하기 전에 페이지를 넘김
        composer.doc.add_page_break()

        # ------ 여기까진 문제 없음 ----------
        
        current_sub_list = files_list[sub_list_index]

        # 이 서브 리스트에 속한 파일들을 같은 페이지에 순차적으로 병합
        # 첫 번째 파일
        first_doc_name = current_sub_list[0]

        doc_temp = Document_compose(os.path.join(full_dir, first_doc_name))

        ### ----------------- 문제 번호 reset -----------------
        if is_reset_num:
            if doc_temp.paragraphs:
                paragraph = doc_temp.paragraphs[0]
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # ✅ 왼쪽 정렬

                if paragraph.runs:
                    paragraph.clear()  # ✅ 기존 번호 삭제
                    new_run = paragraph.add_run(f"{reset_num_count}. ")  # ✅ 새로운 번호 추가
                    new_run.bold = True  # ✅ Bold 적용
                    new_run.font.size = Pt(12)  # ✅ 글자 크기 설정

            reset_num_count += 1  # ✅ 문제 번호 증가
            ### ----------------- 문제 번호 reset -----------------

        composer.append(doc_temp)

        # 두 번째 파일부터는 같은 페이지에서 문단 띄우고 병합
        if len(current_sub_list) > 1:
            for file_name in current_sub_list[1:]:
                composer.doc.add_paragraph()
                doc_temp = Document_compose(os.path.join(full_dir, file_name))

                ### ----------------- 문제 번호 reset -----------------
                if is_reset_num:
                    if doc_temp.paragraphs:
                        paragraph = doc_temp.paragraphs[0]
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # ✅ 왼쪽 정렬

                        if paragraph.runs:
                            paragraph.clear()  # ✅ 기존 번호 삭제
                            new_run = paragraph.add_run(f"{reset_num_count}. ")  # ✅ 새로운 번호 추가
                            new_run.bold = True  # ✅ Bold 적용
                            new_run.font.size = Pt(12)  # ✅ 글자 크기 설정

                    reset_num_count += 1  # ✅ 문제 번호 증가
                ### ----------------- 문제 번호 reset -----------------

                composer.append(doc_temp)

    ### 임시 경로
    # 3) 결과 파일 저장
    
    folder_path = os.path.join(os.getcwd(), "outputs")  # 상대경로 (현재 작업 디렉토리 기준)

    if not os.path.exists(folder_path):  # 폴더 존재 여부 확인
        os.makedirs(folder_path)  # 폴더 생성
    
    composer.save(os.path.join(os.getcwd(), f"outputs/{output_filename}.docx")) # .exe 경로
    # composer.save(r"D:\YearDreamSchool-D\python_projects\msword_pjt\outputs/final_test.docx") # .py 경로







# def combine_all_asnwer_docx(full_dir, files_list, output_filename, is_reset_num):
#     """
#     주어진 file들의 list를 토대로 모든 문서를 결합.

#     full_dir        : 파일들이 위치한 폴더의 절대경로
#     files_list      : 병합할 파일들의 2차원 리스트
#     output_filename : 유저가 입력한 최종 출력물의 이름
#     is_reset_num    : 문제 번호 초기화 여부
#     """
#     # 1) 첫 번째 서브 리스트의 첫 번째 파일을 '마스터 문서'로 설정
#     master = Document_compose(os.path.join(full_dir, files_list[0][0]))
#     composer = Composer(master)

    
#     # 문제 번호 reset을 위한 변수
#     reset_num_count = 1

#     ### ----------------- 문제 번호 reset -----------------
#     if is_reset_num:
#         if master.paragraphs:
#             paragraph = master.paragraphs[0]
#             paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # ✅ 왼쪽 정렬

#             if paragraph.runs:
#                 paragraph.clear()  # ✅ 기존 번호 삭제
#                 new_run = paragraph.add_run(f"{reset_num_count}. ")  # ✅ 새로운 번호 추가
#                 new_run.bold = True  # ✅ Bold 적용
#                 new_run.font.size = Pt(12)  # ✅ 글자 크기 설정

#             reset_num_count += 1  # ✅ 문제 번호 증가
#     ### ----------------- 문제 번호 reset -----------------

#     # 만약 첫 번째 서브 리스트에 파일이 여러 개라면, 나머지 파일을 같은 페이지에 순차적으로 추가
#     first_sub_list = files_list[0]
#     if len(first_sub_list) > 1:
#         # 두 번째 파일부터 병합
#         for file_name in first_sub_list[1:]:
#             composer.doc.add_paragraph()  # 한 줄 띄우고
#             # doc_temp = Document_compose(os.path.join(full_dir, f"{file_name}.docx"))
#             doc_temp = Document_compose(os.path.join(full_dir, file_name))
            
#             ### ----------------- 문제 번호 reset -----------------
#             if is_reset_num:
#                 if doc_temp.paragraphs:
#                     paragraph = doc_temp.paragraphs[0]
#                     paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # ✅ 왼쪽 정렬

#                     if paragraph.runs:
#                         paragraph.clear()  # ✅ 기존 번호 삭제
#                         new_run = paragraph.add_run(f"{reset_num_count}. ")  # ✅ 새로운 번호 추가
#                         new_run.bold = True  # ✅ Bold 적용
#                         new_run.font.size = Pt(12)  # ✅ 글자 크기 설정

#                     reset_num_count += 1  # ✅ 문제 번호 증가
#             ### ----------------- 문제 번호 reset -----------------

#             composer.append(doc_temp)

   

#     # 2) 나머지 서브 리스트(1번째 인덱스부터 끝까지) 처리
#     for sub_list_index in range(1, len(files_list)):
#         # 새로운 서브 리스트를 삽입하기 전에 페이지를 넘김
#         composer.doc.add_page_break()
        
#         current_sub_list = files_list[sub_list_index]

#         # 이 서브 리스트에 속한 파일들을 같은 페이지에 순차적으로 병합
#         # 첫 번째 파일
#         first_doc_name = current_sub_list[0]

#         doc_temp = Document_compose(os.path.join(full_dir, first_doc_name))

#         ### ----------------- 문제 번호 reset -----------------
#         if is_reset_num:
#             if doc_temp.paragraphs:
#                 paragraph = doc_temp.paragraphs[0]
#                 paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # ✅ 왼쪽 정렬

#                 if paragraph.runs:
#                     paragraph.clear()  # ✅ 기존 번호 삭제
#                     new_run = paragraph.add_run(f"{reset_num_count}. ")  # ✅ 새로운 번호 추가
#                     new_run.bold = True  # ✅ Bold 적용
#                     new_run.font.size = Pt(12)  # ✅ 글자 크기 설정

#                 reset_num_count += 1  # ✅ 문제 번호 증가
#             ### ----------------- 문제 번호 reset -----------------

#         composer.append(doc_temp)

#         # 두 번째 파일부터는 같은 페이지에서 문단 띄우고 병합
#         if len(current_sub_list) > 1:
#             for file_name in current_sub_list[1:]:
#                 composer.doc.add_paragraph()
#                 doc_temp = Document_compose(os.path.join(full_dir, file_name))

#                 ### ----------------- 문제 번호 reset -----------------
#                 if is_reset_num:
#                     if master.paragraphs:
#                         paragraph = master.paragraphs[0]
#                         paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # ✅ 왼쪽 정렬

#                         if paragraph.runs:
#                             paragraph.clear()  # ✅ 기존 번호 삭제
#                             new_run = paragraph.add_run(f"{reset_num_count}. ")  # ✅ 새로운 번호 추가
#                             new_run.bold = True  # ✅ Bold 적용
#                             new_run.font.size = Pt(12)  # ✅ 글자 크기 설정

#                         reset_num_count += 1  # ✅ 문제 번호 증가
#                 ### ----------------- 문제 번호 reset -----------------

#                 composer.append(doc_temp)

def combine_all_answer_docx(full_dir, files_list, output_filename, is_reset_num):
    """
    주어진 file들의 list를 토대로 4개씩 같은 페이지에 문서를 결합.

    full_dir        : 파일들이 위치한 폴더의 절대경로
    files_list      : 병합할 파일들의 리스트
    output_filename : 유저가 입력한 최종 출력물의 이름
    is_reset_num    : 문제 번호 초기화 여부
    """
    if not files_list:
        print("파일 목록이 비어 있습니다.")
        return
    
    def reset_number(doc, number):
        """문서의 첫 번째 문단을 문제 번호로 리셋"""
        if doc.paragraphs:
            paragraph = doc.paragraphs[0]
            paragraph.alignment = 0  # 왼쪽 정렬
            if paragraph.runs:
                paragraph.clear()
                new_run = paragraph.add_run(f"{number}. ")
                new_run.bold = True
                new_run.font.size = Pt(12)
    
    # 첫 번째 파일을 기준 문서로 설정
    master = Document_compose(os.path.join(full_dir, files_list[0]))
    composer = Composer(master)

    # 문제 번호 초기화 변수
    reset_num_count = 1

    if is_reset_num:
        reset_number(master, reset_num_count)
        reset_num_count += 1

    # 4개씩 같은 페이지에 넣고, 이후 페이지 넘기기
    for i, file_name in enumerate(files_list[1:], start=1):  # 첫 파일 제외
        doc_temp = Document_compose(os.path.join(full_dir, file_name))

        # 문제 번호 리셋이 활성화된 경우 번호 업데이트
        if is_reset_num:
            reset_number(doc_temp, reset_num_count)
            reset_num_count += 1

        ### 4개씩 같은 페이지에 추가, 이후 페이지 넘기기
        # if i % 5 == 0:
        #     composer.doc.add_page_break()  # 4개마다 새 페이지 추가
        composer.doc.add_paragraph() # 한 줄 띄우고
        composer.append(doc_temp)

    ### 임시 경로
    # 3) 결과 파일 저장
    
    folder_path = os.path.join(os.getcwd(), "outputs")  # 상대경로 (현재 작업 디렉토리 기준)

    if not os.path.exists(folder_path):  # 폴더 존재 여부 확인
        os.makedirs(folder_path)  # 폴더 생성
    
    composer.save(os.path.join(os.getcwd(), f"outputs/{output_filename}.docx")) # .exe 경로
    # composer.save(r"D:\YearDreamSchool-D\python_projects\msword_pjt\outputs/final_test.docx") # .py 경로



def combine_all_docx_one_by_one(full_dir, files_list, output_filename, is_reset_num):
    """
    주어진 file들을 각각 다른 페이지에 넣어 모두 병합.

    full_dir       : 파일 전까지의 전체 절대경로로
    files_list     : 병합할 파일들의 list
    utput_filename : 유저가 입력한 최종 출력물의 이름
    is_reset_num   : 문제 번호 초기화 여부
    """
    # master문서(첫 시작 문서) 설정
    master = Document_compose(os.path.join(full_dir, f"{files_list[0]}"))
    # master문서를 composer에 할당
    composer = Composer(master)

    ### ----------------- 문제 번호 reset -----------------
    reset_num_count = 1
    
    if is_reset_num:
        if master.paragraphs:
            paragraph = master.paragraphs[0]
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # ✅ 왼쪽 정렬

            if paragraph.runs:
                paragraph.clear()  # ✅ 기존 번호 삭제
                new_run = paragraph.add_run(f"{reset_num_count}. ")  # ✅ 새로운 번호 추가
                new_run.bold = True  # ✅ Bold 적용
                new_run.font.size = Pt(12)  # ✅ 글자 크기 설정

            reset_num_count += 1  # ✅ 문제 번호 증가
    ### ----------------- 문제 번호 reset -----------------

    for i in range(1, len(files_list)):
        # composer.doc.add_paragraph() # 한 칸 띄어쓰기 추가
        composer.doc.add_page_break() # 다음 페이지로
        composer.doc.add_paragraph()  # 한 줄 띄우고
        doc_temp = Document_compose(os.path.join(full_dir, f"{files_list[i]}")) # 다음 문서 가져오기

        ### ----------------- 문제 번호 reset -----------------
        # ✅ 첫 번째 문단 수정 (번호 변경 + Bold + 왼쪽 정렬)
        if is_reset_num:
            if doc_temp.paragraphs:
                paragraph = doc_temp.paragraphs[0]
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # ✅ 왼쪽 정렬

                if paragraph.runs:
                    paragraph.clear()  # ✅ 기존 번호 삭제
                    new_run = paragraph.add_run(f"{reset_num_count}. ")  # ✅ 새로운 번호 추가
                    new_run.bold = True  # ✅ Bold 적용
                    new_run.font.size = Pt(12)  # ✅ 글자 크기 설정
        ### ----------------- 문제 번호 reset -----------------

        reset_num_count += 1
        composer.append(doc_temp) # 가져온 temp문서 master문서에 append



    # 결과 파일 저장
    folder_path = os.path.join(os.getcwd(), "outputs")  # 상대경로 (현재 작업 디렉토리 기준)

    if not os.path.exists(folder_path):  # 폴더 존재 여부 확인
        os.makedirs(folder_path)  # 폴더 생성
    
    composer.save(os.path.join(os.getcwd(), f"outputs/{output_filename}.docx")) # .exe 경로
    # composer.save(r"D:\YearDreamSchool-D\python_projects\msword_pjt\outputs/final_test.docx") # .py 경로




def combine_all_docx_seamless(full_dir, files_list, output_filename, is_reset_num):
    """
    주어진 file들을 각각 다른 페이지에 넣어 모두 병합.

    full_dir       : 파일 전까지의 전체 절대경로로
    files_list     : 병합할 파일들의 list
    utput_filename : 유저가 입력한 최종 출력물의 이름
    is_reset_num   : 문제 번호 초기화 여부
    """
    # master문서(첫 시작 문서) 설정
    master = Document_compose(os.path.join(full_dir, f"{files_list[0]}"))
    # master문서를 composer에 할당
    composer = Composer(master)

    ### ----------------- 문제 번호 reset -----------------
    reset_num_count = 1
    
    if is_reset_num:
        if master.paragraphs:
            paragraph = master.paragraphs[0]
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # ✅ 왼쪽 정렬

            if paragraph.runs:
                paragraph.clear()  # ✅ 기존 번호 삭제
                new_run = paragraph.add_run(f"{reset_num_count}. ")  # ✅ 새로운 번호 추가
                new_run.bold = True  # ✅ Bold 적용
                new_run.font.size = Pt(12)  # ✅ 글자 크기 설정

            reset_num_count += 1  # ✅ 문제 번호 증가
    ### ----------------- 문제 번호 reset -----------------

    for i in range(1, len(files_list)):
        composer.doc.add_paragraph() # 한 칸 띄어쓰기 추가
        # composer.doc.add_page_break() # 다음 페이지로
        doc_temp = Document_compose(os.path.join(full_dir, f"{files_list[i]}")) # 다음 문서 가져오기

        ### ----------------- 문제 번호 reset -----------------
        # ✅ 첫 번째 문단 수정 (번호 변경 + Bold + 왼쪽 정렬)
        if is_reset_num:
            if doc_temp.paragraphs:
                paragraph = doc_temp.paragraphs[0]
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # ✅ 왼쪽 정렬

                if paragraph.runs:
                    paragraph.clear()  # ✅ 기존 번호 삭제
                    new_run = paragraph.add_run(f"{reset_num_count}. ")  # ✅ 새로운 번호 추가
                    new_run.bold = True  # ✅ Bold 적용
                    new_run.font.size = Pt(12)  # ✅ 글자 크기 설정
        ### ----------------- 문제 번호 reset -----------------

        reset_num_count += 1
        composer.append(doc_temp) # 가져온 temp문서 master문서에 append



    # 결과 파일 저장
    folder_path = os.path.join(os.getcwd(), "outputs")  # 상대경로 (현재 작업 디렉토리 기준)

    if not os.path.exists(folder_path):  # 폴더 존재 여부 확인
        os.makedirs(folder_path)  # 폴더 생성
    
    composer.save(os.path.join(os.getcwd(), f"outputs/{output_filename}.docx")) # .exe 경로
    # composer.save(r"D:\YearDreamSchool-D\python_projects\msword_pjt\outputs/final_test.docx") # .py 경로